import streamlit as st
import pdfplumber
import re
from google import genai
from google.genai import types
from docx import Document
from io import BytesIO
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from huggingface_hub import InferenceClient
import os  # <-- ADD THIS
import html

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# --- FIX "INVALID PORT" NETWORK ERROR ---
for proxy_var in ['http_proxy', 'https_proxy', 'HTTP_PROXY', 'HTTPS_PROXY']:
    if proxy_var in os.environ:
        del os.environ[proxy_var]

st.session_state.setdefault("download_ready", False)
# --- NEW SESSION STATE VARIABLES ---
st.session_state.setdefault("doc_generated", False)
st.session_state.setdefault("original_letter", "")
st.session_state.setdefault("edited_letter", "")
st.session_state.setdefault("file_prefix", "") # <-- ADD THIS LINE
# -----------------------------------

# -----------------------------
# CONFIGURATION (SECRETS)
# -----------------------------
GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", "").strip()
HF_TOKEN = st.secrets.get("HF_TOKEN", "").strip()
HF_MODEL = st.secrets.get("HF_MODEL", "meta-llama/Meta-Llama-3-8B-Instruct").strip()


if not GEMINI_API_KEY and not HF_TOKEN:
    st.error("Missing GEMINI_API_KEY and HF_TOKEN in Streamlit secrets.")
    st.stop()


import platform
import shutil

# Cross-platform Tesseract configuration
if platform.system() == "Windows":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
else:
    # On Linux (Streamlit Cloud), find tesseract in PATH
    tess_path = shutil.which("tesseract")
    if tess_path:
        pytesseract.pytesseract.tesseract_cmd = tess_path
    else:
        print("⚠️ Tesseract not found on system PATH. OCR fallback may fail.")

# -----------------------------
# STREAMLIT FRONT END SETUP
# -----------------------------
st.set_page_config(page_title="BCBS Justification for IDR", layout="centered")
st.title("🏥 BCBS Justification for IDR Generator")
st.write("Upload EOB PDF, MRN PDF, and Prompt TXT file to generate a formatted BCBS justification document.")

# -----------------------------
# FILE UPLOADS
# -----------------------------
# -----------------------------
# -----------------------------
# FILE UPLOADS & INPUTS
# -----------------------------
multiple_eobs = st.toggle("Enable Multiple EOB Upload", value=True)
eob_files = st.file_uploader("📄 Upload EOB PDFs", type=["pdf"], accept_multiple_files=multiple_eobs)
mrn_file = st.file_uploader("🧾 Upload MRN PDF", type=["pdf"])
target_cpt_code_input = st.text_input("🔢 Enter Target CPT Code (e.g., 99283, 99284)", value="")

# -----------------------------
# HELPER FUNCTIONS
# -----------------------------
def extract_text_from_pdf(uploaded_file):
    """Extract text from both scanned and digital PDFs"""
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except Exception:
        pass
    uploaded_file.seek(0)
    images = convert_from_bytes(uploaded_file.read())
    for img in images:
        text += pytesseract.image_to_string(img) + "\n"
    return text


def find_field(patterns, text, label):
    """Try multiple regex patterns until a match is found"""
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return f"{label} not found"


def extract_fields(eob_text):
    """Extract claim fields robustly from EOB text"""
    # ---------------- DATE OF SERVICE ----------------
    date_patterns = [
        r"Date.?of.?Service[s]?:?\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Service Date[s]?:?\s*(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Service Dates?\s*[:\-–\s]+(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Date Range[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})",
        r"Dates? of Service[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})"
    ]
    date_of_service = find_field(date_patterns, eob_text, "Date")

    # ---------------- HCPCS / CPT CODES ----------------
    # ---------------- HCPCS / CPT CODES ----------------
    # ---------------- HCPCS / CPT CODES ----------------
    lines = eob_text.splitlines()
    ranked_codes = []

    # Emergency E/M codes pattern
    emergency_pattern = r"99(28[1-5]|29[1-2])"

    for line in lines:
        code_match = re.search(r"(?:HCPCS|CPT|^|\s)([012789]\d{4}[A-Z]?|0\d{3}[A-Z])(?:\s|$)", line)
        if not code_match:
            continue

        code = code_match.group(1)

        # Ignore mixed alphanumeric strings
        if re.search(r"[A-Za-z]{2,}\d{4,}", line):
            continue

        base_code = code

        # Automatically add -25 to ALL emergency codes
        if re.match(emergency_pattern, base_code):
            code = f"{base_code}-25"

        if code not in ranked_codes:
            ranked_codes.append(code)


    # ---------------- DRG CODE ----------------
    drg_patterns = [
        r"DRG\s*[:#-]?\s*([0-9]{2,4})",
        r"DRG\s*Code\s*[:#-]?\s*([0-9]{2,4})",
        r"Diagnosis\s*Related\s*Group[^0-9]*([0-9]{2,4})",
        r"MS[-\s]*DRG[^0-9]*([0-9]{2,4})",
        r"DRG[^0-9]*([0-9]{2,4})",
        r"RelatedGroup[^0-9]*([0-9]{2,4})"
    ]
    drg_code = find_field(drg_patterns, eob_text, "DRG Code")

# ---------------- BILLING PROVIDER ----------------
    # Updated to capture multi-line names and table formats in PDFs
    billing_patterns = [
        # Multi-line/Table patterns with broader anchors and positive lookahead
        # Prioritize Payee Name as it often contains the full legal name in table formats
        r"Payee Name[\"\, \s\n]*([\s\S]*?)(?=\"?\s*(?:Check Date|Line Level|Prior Notification|NPI|Address|Billing Provider|Rendering Provider))",
        r"Billing Provider Name[\"\, \s\n]*([\s\S]*?)(?=\"?\s*(?:Billing Provider NPI|Rendering Provider|Payee Name|Address|NPI|City|State|Zip|Phone|Check Date|Tax ID|Provider|Control Number))",
        r"Billing Provider[\"\, \s\n]*([\s\S]*?)(?=\"?\s*(?:Billing Provider NPI|Rendering Provider|Payee Name|Address|NPI|City|State|Zip))",
        r"Provider Name[\"\, \s\n]*([\s\S]*?)(?=\"?\s*(?:Billing Provider NPI|Rendering Provider|Payee Name|Address|NPI|City|State|Zip))",
        # Fallback to greedy character set (now includes quotes and commas for cleaning)
        r"Billing Provider[:\s]*([A-Za-z0-9\s.,&'\"\-]+)",
        r"Provider Name[:\s]*([A-Za-z0-9\s.,&'\"\-]+)"
    ]

    billing_provider = find_field(billing_patterns, eob_text, "Billing Provider")

    # Clean up the extracted string (remove PDF table artifacts like quotes, commas, and newlines)
    billing_provider = re.sub(r'[\"\,\n]+', ' ', billing_provider)
    
    # Remove interleaved labels that might have been caught by the broad capture
    billing_provider = re.sub(r'\b(?:Payee Name|Billing Provider Name|Provider Name|Rendering Provider Name)\b', ' ', billing_provider, flags=re.IGNORECASE)
    
    billing_provider = re.sub(r'\s+', ' ', billing_provider).strip()

    # Remove trailing noise just in case it over-captures
    billing_provider = re.sub(
        r"\s*(NPI.*|Other Carrier.*|Rendering Provider.*|Check Date.*|Address.*|City.*|State.*|Zip.*|Tax ID.*)$",
        "",
        billing_provider,
        flags=re.IGNORECASE,
    ).strip()

    # Deduplicate repeated name suffixes (common in table wraps like "CENTER LLC CENTER LLC")
    words = billing_provider.split()
    if len(words) >= 4:
        for i in range(1, len(words) // 2 + 1):
            if words[-i:] == words[-2*i:-i]:
                billing_provider = " ".join(words[:-i])
                break

    return date_of_service, ranked_codes, drg_code, billing_provider

import time

def _is_quota_or_rate_limit_error(e: Exception) -> bool:
    msg = str(e).lower()
    return ("429" in msg) or ("quota" in msg) or ("rate limit" in msg) or ("resource_exhausted" in msg)


def _truncate_for_llm(text: str, max_chars: int = 20000) -> str:
    text = text or ""
    if len(text) <= max_chars:
        return text
    # keep head + tail so you don’t lose discharge/impression sections
    head = int(max_chars * 0.7)
    tail = max_chars - head
    return text[:head] + "\n\n[TRUNCATED]\n\n" + text[-tail:]


def _hf_generate(prompt: str, hf_token: str, hf_model: str) -> str:
    if not hf_token:
        raise RuntimeError("HF_TOKEN is missing in Streamlit secrets.")

    prompt = _truncate_for_llm(prompt, max_chars=20000)

    client = InferenceClient(token=hf_token)

    try:
        resp = client.chat.completions.create(
            model=hf_model,
            messages=[
                {"role": "system", "content": "Return a structured MRN summary with **bold** headings and line breaks. Clinical, concise."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=900,
            temperature=0.2,
        )
        return resp.choices[0].message.content.strip()

    except Exception as e:
        # Try to extract HF router JSON error details
        detail = ""
        resp_obj = getattr(e, "response", None)
        if resp_obj is not None:
            try:
                detail = resp_obj.json()
            except Exception:
                try:
                    detail = resp_obj.text
                except Exception:
                    detail = ""

        raise RuntimeError(f"HF router error: {e}\nDETAIL: {detail}")




def generate_mrn_summary(mrn_text, target_cpt_code):
    """Gemini primary -> Gemini fallback sequence -> Hugging Face fallback."""
    
    # HARDCODED PROMPT WITH DYNAMIC CPT CODE INJECTION
    prompt_text = f"""Based on the uploaded medical record, please generate a formal "Clinical Summary and CPT Code Justification" for medical billing purposes.
The output must follow the exact structure and professional tone below.
Focus on medical necessity, severity of illness, and the complexity of medical decision-making (MDM). 

Structure the response as follows:

**I. Clinical Summary and Severity of Illness**
* Start with the patient demographics (age/sex), relevant history, and mode of arrival (don't mention patient name).
* State the Acuity Level and describe the patient's condition upon arrival (symptoms, pain scores, duration).
* Corroborate severity with objective data: Vital signs (highlight abnormalities like tachycardia, hypertension, hypoxia, etc.), Physical Exam findings, and key Lab/Imaging results (cite specific values).
* do not mention Low complexity, Stable, Negative imaging, Acuity Less Urgent.

**II. Impact of Treatment and Medical Necessity**
* Explain why the treatments (IV fluids, meds, imaging) were medically necessary to prevent deterioration or rule out life-threatening conditions.
* Describe the specific interventions (e.g., "1000 mL Normal Saline," specific meds) and their direct impact on the patient (e.g., "reduced pain score from X to Y," "stabilized vitals").

**III. CPT Code Justification**
* Justify the following specific codes based on the evidence in the chart: **{target_cpt_code}**
    * For each code listed, provide a distinct bullet point explaining the specific medical necessity and chart evidence that supports it. Do not use bold text for the bullet points.

**IV. Conclusion**
* Conclude with a single, high-impact paragraph starting with exactly this phrase: "The medical record clearly demonstrates that this encounter was not a routine evaluation of [Chief Complaint]."
* and systemic viral symptoms necessitated the intensive resource utilization reflected in CPT codes **{target_cpt_code}**.

**Instructions:**
* Use professional, billing-focused language (e.g., "profound physiological instability," "corroborated by," "substantiated by").
* Be specific with data points from the file (e.g., Creatinine 1.4 mg/dL, WBC 15.1, etc.).
* Format exactly like the example provided in your training data."""

    mrn_text = _truncate_for_llm(mrn_text, max_chars=18000)
    combined_prompt = (
        f"{prompt_text}\n\n---\n\n{mrn_text}\n\n"
        "Format output using **bold** section headings and line breaks for clarity. Do not use bold text inside bullet points."
    )

    last_gemini_error = None

    if GEMINI_API_KEY:
        try:
            # Add http_options to force REST connection
            client = genai.Client(
                api_key=GEMINI_API_KEY,
                http_options={'api_version': 'v1beta'}
            )
            gemini_models = [
                "gemini-2.5-flash",
                "gemini-2.5-flash-lite",
                "gemini-3-flash-preview",
                "gemini-3.5-flash"      
            ]

            for i, model_name in enumerate(gemini_models):
                try:
                    print(f"Attempting Gemini model: {model_name}...")
                    response = client.models.generate_content(
                        model=model_name,
                        contents=combined_prompt
                    )
                    st.session_state["llm_used"] = model_name
                    return format_summary_bullets(response.text.strip())
                except Exception as model_error:
                    last_gemini_error = model_error
                    if i < len(gemini_models) - 1:
                        if _is_quota_or_rate_limit_error(model_error):
                            time.sleep(2)
                        continue 
                    else:
                        raise model_error 

        except Exception as e2:
            st.warning(f"⚠️ Gemini failed, falling back to HF. Reason: {e2}")
            last_gemini_error = e2

    # HF fallback
    try:
        out = _hf_generate(combined_prompt, HF_TOKEN, HF_MODEL)
        st.session_state["llm_used"] = f"hf:{HF_MODEL}"
        return format_summary_bullets(out)
    except Exception as e3:
        msg = str(e3).lower()
        if "403" in msg and "inference providers" in msg:
            raise RuntimeError("HF 403: Inference Providers permission needed.")
        if "not a chat model" in msg or "model_not_supported" in msg:
            raise RuntimeError(f"HF model/endpoint mismatch for '{HF_MODEL}'.")
        raise RuntimeError(f"HF failed: {e3}")






def generate_bcbs_justification_letter(date, hcpcs, drg, billing_provider, mrn_summary):
    """Build the complete letter text with variables inserted."""
    
    # --- SORT HCPCS CODES (Emergency codes first) ---
    emergency_codes = []
    other_codes = []
    for c in hcpcs:
        base_code = re.sub(r"-.*", "", c)
        if re.match(r"99(28[1-5]|29[1-2])", base_code):
            emergency_codes.append(c) # Append the full code (e.g. 99285-25)
        else:
            other_codes.append(c)
            
    # Combine lists: Emergency codes first, then everything else
    sorted_hcpcs = emergency_codes + other_codes
    hcpcs_list = ', '.join(sorted_hcpcs) if sorted_hcpcs else 'N/A'
    
    # Grab just the base emergency code for the modifier 25 sentence
    emergency_code_text = re.sub(r"-.*", "", emergency_codes[0]) if emergency_codes else "99284"

    return f"""
This letter is submitted in support of our Independent Dispute Resolution (IDR) request under the No Surprises Act (NSA). We are challenging the reimbursement amount determined by BCBS for the emergency services rendered on **{date}**. The payment issued by BCBS does not adequately reflect the level of care provided, nor does it comply with NSA transparency requirements.

We firmly assert that a higher reimbursement is justified based on the significant medical complexity of the case and in accordance with the payment determination criteria outlined in **45 CFR §149.510(c)(4)(iii)**. This includes, but is not limited to, the acuity of the patient’s condition, the scope of services rendered, and the qualifications and experience of the attending provider.

### QPA Transparency Failure & Arbitrary Methodology
BCBS’s claims their QPA was calculated using internal, fee-for-service median contracted rates from 2020–2021 and that it excludes bonuses, shared risk, and information derived from databases. However, this explanation fails to meet the disclosure standards under 45 CFR §149.140(a)(12) and CMS NSA FAQ #12, which require the QPA to be provided on a per-service basis and not as a flat amount across unrelated CPT’s.
Furthermore, BCBS does not disclose:
•	The actual median rate or how it was derived
•	Which specialties or contracts were used
•	Whether their data reflects ghost rates, $1 floor rates, or stale agreements
These omissions contradict the purpose of the NSA’s transparency goals and materially impair providers' ability to evaluate fairness and negotiate in good faith. In effect, BCBS is asking IDR entities to accept a "black box" QPA with no line-item disclosure, no clinical justification, and no meaningful accountability.
This approach also disproportionately harms out-of-network emergency providers, especially those serving underserved populations and operating independently of hospital systems. It reinforces the need for a fair, case-specific evaluation of the actual services rendered, which far exceeds the clinical complexity of what BCBS’s rate represents.

###  Improper DRG Classification and Non-Compliant Adjudication of Outpatient Emergency Claim
A thorough assessment of the Explanation of Benefits (EOB) provided by Blue Cross Blue Shield (BCBS) reveals major systemic errors in claim adjudication and misclassification. The claim was properly submitted using multiple, separate CPT/HCPCS codes—**{hcpcs_list}**—each reflecting distinct, medically necessary procedures performed during the emergency department encounter. Modifier 25 was correctly appended to CPT **{emergency_code_text}**, denoting a significant, separately identifiable evaluation and management (E/M) service delivered in addition to diagnostic testing, as recognized by CMS's National Correct Coding Initiative (NCCI) policy.
Nevertheless, the BCBS EOB lists a “Hospital Payment Indicator: R – Case Rate” and assigns Diagnosis Related Group (DRG) **{drg}** with a zero DRG weight (0.00000), indicating the payer’s system improperly converted an outpatient freestanding emergency department (FSED) claim into an inpatient, DRG-based payment methodology. This reclassification is both factually incorrect and in violation of federal billing and adjudication requirements. DRG payment models are explicitly reserved for inpatient hospital stays and are not permitted as a basis for adjudicating outpatient emergency claims billed under CPT/HCPCS coding protocols.
Applying DRG automation to an FSED claim misrepresents the nature of the service, the facility type, and the context of care delivery, resulting in an unsupported case-rate payment that fails to consider the submitted codes and the true scope of services rendered. This process does not satisfy the legal requirements of 45 C.F.R. § 149.510(c)(4)(iii), which obligate payers to evaluate payment based on the provider’s experience, facility type, service scope, and patient acuity.
For these reasons, it is imperative that the adjudication and Independent Dispute Resolution (IDR) review privilege only the original CPT/HCPCS codes submitted, accurately representing the outpatient emergency care delivered. Reimbursement must be recalculated according to NSA regulations to guarantee precise, transparent, and equitable payment in line with the intent of the No Surprises Act.

### Patient Acuity & Complexity of Care
{mrn_summary}

### Training, Experience & Quality Measures
**{billing_provider}** operates as a 24/7 Freestanding Emergency Department (FSED), staffed exclusively by board-certified emergency medicine physicians and highly trained nursing professionals dedicated to delivering care that meets or exceeds nationally recognized benchmarks for clinical accuracy, quality, and patient safety. Our commitment to excellence is demonstrated through accreditation by The Joint Commission (Gold Seal of Approval), COLA’s Seal of Quality in Healthcare, and the Center for Improvement in Healthcare Quality (CIHQ)—each reflecting rigorous national compliance standards for safety, quality, and patient outcomes.
We have made significant capital and operational investments in advanced diagnostic and treatment technologies, including multi-slice CT scanners, digital radiography, and on-site laboratory services, enabling our team to deliver hospital-level emergency care efficiently while reducing the overcrowding burden typical of traditional hospital ERs. FSEDs in Texas are highly regulated under state licensure laws, undergo continuous inspections, and are bound by EMTALA-comparable obligations to ensure all patients receive appropriate emergency care regardless of insurance status or ability to pay. Independent research consistently demonstrates that Texas FSEDs deliver timely, efficient, and medically necessary care comparable to—if not exceeding—hospital-based emergency departments, particularly regarding patient throughput and satisfaction.
These operational requirements and the higher-acuity case mix we routinely manage are not represented in the payer’s QPA dataset and therefore warrant an upward deviation. The Qualified Payment Amount (QPA) is derived from median in-network rates that often blend hospital outpatient and urgent-care data—entities that do not share our 24/7 readiness, staffing ratios, or advanced clinical scope. Moreover, as a community-based emergency care provider, our case mix includes a broad range of high-acuity and after-hours presentations that cannot safely be managed by “ordinary providers” or urgent care facilities.
Accordingly, under 45 C.F.R. § 149.510(c)(4)(iii)(C), the certified IDR entity must give substantial weight to the provider’s scope of services, case mix, and clinical capabilities when determining the appropriate out-of-network rate. Our facility’s distinct operational and clinical profile materially differentiates us from the entities included in BCBS’s QPA calculation and establishes that the payer’s presumptive rate fails to capture the true cost and complexity of emergency care delivered in this setting.
This evidence affirms the critical role of **{billing_provider}** in providing high-acuity, hospital-level emergency services that uphold national quality standards and advance the protections intended under the No Surprises Act, ensuring that reimbursement determinations reflect both fairness and the indispensable public-health function of FSEDs

### Teaching Status, Case Mix, & Scope of Services
As a freestanding emergency center, **{billing_provider}** provides comprehensive emergency care across a broad and complex case mix. Our facility manages adult and pediatric emergencies, including critical conditions such as stroke and trauma stabilization, acute abdominal, cardiac complaints, and mental health evaluations. This diverse case mix reflects the high acuity and clinical complexity our physicians address daily.
Our scope of services is extensive and aligns with the full continuum of emergency care, including but not limited to:
•	Advanced diagnostic imaging (multi-slice CT scans, digital radiography)
•	Comprehensive laboratory diagnostics (CBC, BMP, LFTs, cardiac biomarkers)
•	Emergency medication administration, including critical care drugs
•	Intravenous fluid therapy and resuscitation
•	Short-term patient observation and stabilization
•	Coordinated emergency discharge planning and seamless transitions to higher levels of care or outpatient follow-up
Our board-certified emergency physicians and highly trained clinical staff maintain readiness to manage high-acuity patients around the clock, ensuring rapid response and quality outcomes. This case exemplifies the breadth and depth of our clinical capabilities and the critical nature of services provided by a Texas FSED.

### Market Share Considerations
Our facility serves as a vital safety net within the regional healthcare ecosystem, particularly during high-demand periods when hospital emergency departments exceed capacity. As a nonparticipating provider, our patient volume is inherently limited compared to Blue Cross Blue Shield Association (BCBS) expansive regional footprint. BCBS’s dominant market share creates a significant power imbalance, allowing them to exert disproportionate leverage over the calculation of Qualified Payment Amounts (QPAs). Consequently, these benchmarks are artificially depressed and fail to account for the actual costs and clinical complexities associated with emergency care.

### Negotiations and Good Faith Efforts
Acceptance of the plan’s proposed in-network reimbursement rates would not be operationally sustainable for a licensed freestanding emergency medical care facility. The proposed rates are materially insufficient to support the fixed operational and clinical costs associated with maintaining continuous 24/7 emergency readiness, including coverage by board-certified emergency medicine physicians, licensed nursing staff, and essential ancillary personnel. Operating under such reimbursement levels would undermine the facility’s financial viability and impair its ability to meet the clinical, staffing, and safety obligations inherent to emergency care delivery.
Maintaining non-participating status is therefore not a discretionary business decision, but a necessary measure to ensure that emergency physicians and staff are compensated fairly, regulatory and safety standards are upheld, and the facility remains equipped with advanced diagnostic tools and life-saving medical equipment essential to timely treatment and optimal patient outcomes. The No Surprises Act does not require emergency providers to accept reimbursement terms that compromise their ability to deliver compliant, high-quality emergency services. Accordingly, the plan’s contracted rates—and any Qualified Payment Amount derived from them—do not reasonably reflect the value, complexity, or resource intensity of the emergency services furnished in this case and should not be afforded controlling weight in the IDR determination. 
As a freestanding emergency department, our statutory duty under the Prudent Layperson Standard is to provide emergency medical care to any patient presenting with symptoms of a potential emergency—without regard to insurance status or network participation. We are not permitted to defer or deny care based on contractual considerations. The lack of a network agreement, therefore, reflects not a refusal by the provider to participate, but rather BCBS’s use of market dominance to impose unsustainable contract terms. Under 45 C.F.R. § 149.510(c)(4)(iii)(B), credible evidence of such contracting history requires the certified IDR entity to consider whether the plan’s conduct has distorted the Qualified Payment Amount (QPA). Because BCBS’s self-reported QPA is derived from an in-network dataset shaped by its own suppressed contract rates, the presumption of QPA accuracy should not apply. Accordingly, this factor supports an upward adjustment to the QPA to reflect fair market value for the emergency services rendered.

### Other 
Under 45 C.F.R. § 149.510(c)(4)(iii), the plan’s Qualifying Payment Amount (QPA)—defined as the plan-calculated median in-network rate from 2019, adjusted for inflation—is presumed to represent the appropriate out-of-network (OON) rate. However, this presumption is fundamentally flawed when applied to freestanding emergency departments (FSEDs) due to key methodological and market limitations inherent in the payer’s QPA calculation.
Most FSEDs operate as non-contracted facilities, resulting in minimal or no in-network data from which a valid median contracted rate can be derived. Consequently, payers often substitute unrelated facility data—such as urgent care centers, physician offices, or hospital outpatient departments—to estimate QPAs. This approach violates 45 C.F.R. § 149.140(a)(8), which mandates that QPAs be based on “similar items and services” furnished by providers in the same or comparable specialty and facility type.
Payers frequently calculate and report QPAs based on self-reported, unaudited internal data, lacking external verification and methodological transparency, in breach of the disclosure requirements under 45 C.F.R. § 149.140(a)(12). This opacity obstructs providers’ ability to meaningfully evaluate the accuracy or fairness of proposed payments, thereby impairing the good-faith negotiation process envisioned by the No Surprises Act.
FSEDs incur substantial 24/7 operational and clinical readiness costs equivalent to those of hospital-based emergency departments—such as maintaining board-certified emergency physicians, on-site advanced imaging (CT, ultrasound, digital radiography), full laboratory services, intravenous medication administration, and critical-care stabilization capabilities. These significant fixed and standby expenses are inherently excluded from payer QPA algorithms, which primarily rely on contracted rates for lower-acuity or outpatient care settings.
Therefore, while payers assert their QPA reflects a “fair” OON rate, the underlying data are incomplete, non-comparable, and unverifiable, leading to a QPA that materially misrepresents the true cost and complexity of care delivered by FSEDs. Pursuant to § 149.510(c)(4)(iv), the provider submits credible, facility-specific evidence demonstrating that the QPA diverges materially from the appropriate market rate, including:
– Documentation of the training and board certification of attending emergency physicians;
– Evidence of the acuity and complexity inherent in emergency encounters routinely managed at the facility;
– Description of the scope and availability of emergency diagnostic and treatment resources;
– Verification of the facility’s continuous operational readiness and higher fixed costs; and
– Independent FAIR Health 80th-percentile benchmark data for ZIP code [insert ZIP], corroborating regional commercial rates considerably above the plan’s stated QPA.
None of this evidence relies on prohibited factors under § 149.510(c)(4)(v)—such as billed charges, usual and customary charges, or public payor rates. Instead, these data collectively illustrate that the payer’s QPA fails to capture the genuine cost and complexity of freestanding emergency care. Accordingly, the provider’s submitted rate represents the most accurate, market-reflective reimbursement aligned with the statutory intent and payment-determination framework of the No Surprises Act.

### Conclusion
In summary, the evidence clearly demonstrates that Blue Cross Blue Shield’s payment does not reflect the actual complexity or cost of the emergency services rendered, nor does it comply with the transparency and fairness standards established under the No Surprises Act and 45 C.F.R. §§149.140 and 149.510. By failing to disclose per-service QPAs and relying on internal, non-verifiable methodologies, BCBS has deprived the provider of the ability to assess payment accuracy or negotiate in good faith—contrary to both the spirit and letter of federal law.

**{billing_provider}** has consistently acted in good faith, delivering board-certified, 24/7 emergency care that meets nationally recognized clinical and quality standards. The services in this case were medically necessary, appropriately coded, and supported by full documentation. BCBS’s reliance on a “case-rate” framework misrepresents the nature of outpatient emergency billing and undermines equitable reimbursement practices.
We therefore respectfully request that the certified IDR entity issue a determination in favor of the provider’s offer. Such a decision would uphold the statutory payment-determination factors under 45 C.F.R. §149.510(c)(4)(iii), reinforce transparency in payer conduct, and preserve fair access to emergency medical care within the community.


""".strip()


# -----------------------------
# HELPER FOR BULLET BOLDING
# -----------------------------
def fix_bullet_title_bolding(s):
    """Ensure that for bullet points with a title ending in a colon, ONLY the title part before the colon is bolded."""
    s_trimmed = s.strip()
    bullet_prefix = ""
    if s_trimmed.startswith("•"):
        bullet_prefix = "• "
        content = s_trimmed[1:].strip()
    elif s_trimmed.startswith("* "):
        bullet_prefix = "* "
        content = s_trimmed[2:].strip()
    elif s_trimmed.startswith("- "):
        bullet_prefix = "- "
        content = s_trimmed[2:].strip()
    else:
        return s

    if ":" in content:
        title_part, rest_part = content.split(":", 1)
        clean_title = re.sub(r"\*+", "", title_part).strip()
        clean_rest = re.sub(r"\*+", "", rest_part).strip()
        if clean_rest:
            return f"{bullet_prefix}**{clean_title}:** {clean_rest}"
        else:
            return f"{bullet_prefix}**{clean_title}:**"
    return f"{bullet_prefix}{re.sub(r'\*+', '', content)}"


def format_summary_bullets(summary_text):
    if not summary_text:
        return summary_text
    lines = summary_text.split("\n")
    processed_lines = [fix_bullet_title_bolding(line) for line in lines]
    return "\n".join(processed_lines)


# -----------------------------
# DOCX CREATION FUNCTIONS
# -----------------------------
def parse_bold_segments(paragraph, text):
    """Handle bold (**text**) and remove stray asterisks."""
    text = re.sub(r"\*\*([A-Za-z0-9\s:/()\-]+)\*+", r"**\1**", text)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2].strip())
            run.bold = True
        else:
            paragraph.add_run(part.replace("*", ""))


def add_formatted_paragraph(doc, text):
    """Handle headings, bullets, and normal paragraphs."""
    s = fix_bullet_title_bolding(text.strip())
    if not s:
        doc.add_paragraph("")
        return

    # Headings
    if s.startswith("### ") or re.match(r"^[A-Z][A-Za-z\s&]+:$", s):
        heading = doc.add_paragraph()
        run = heading.add_run(s.replace("### ", "").strip(":").strip())
        run.bold = True
        run.underline = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(13)
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(6)
        return

    # Bulleted lists
    if s.startswith("•") or s.startswith("* "):
        clean_text = s.replace("•", "").replace("* ", "").strip()
        p = doc.add_paragraph(style="List Bullet")
        parse_bold_segments(p, clean_text)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        return

    # Regular text
    p = doc.add_paragraph()
    parse_bold_segments(p, s)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)


def create_docx_with_full_letter(full_letter):
    """Generate final DOCX with Aptos (Body) font."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Aptos (Body)"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos (Body)")
    style.font.size = Pt(12)

    for line in full_letter.split("\n"):
        add_formatted_paragraph(doc, line)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# -----------------------------
# PDF CREATION FUNCTIONS
# -----------------------------
def format_markdown_for_reportlab(text):
    """Convert markdown formatting (**bold**) and special chars to ReportLab XML format."""
    parts = text.split("**")
    formatted_parts = []
    for i, part in enumerate(parts):
        escaped = html.escape(part)
        if i % 2 == 1 and i < len(parts) - 1:  # Inside **...** and properly closed
            formatted_parts.append(f"<b>{escaped}</b>")
        else:
            formatted_parts.append(escaped)
    res = "".join(formatted_parts)
    return res.replace("*", "")


def create_pdf_with_full_letter(full_letter):
    """Generate final PDF with formatting."""
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=14,
        spaceAfter=6
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12.5,
        leading=16,
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )

    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10.5,
        leading=14,
        leftIndent=20,
        spaceAfter=4
    )

    story = []

    for line in full_letter.split("\n"):
        s = fix_bullet_title_bolding(line.strip())
        if not s:
            story.append(Spacer(1, 6))
            continue

        # Headings
        if s.startswith("### ") or re.match(r"^[A-Z][A-Za-z\s&]+:$", s):
            clean_heading = s.replace("### ", "").strip(":").strip()
            formatted_h = format_markdown_for_reportlab(clean_heading)
            story.append(Paragraph(f"<u><b>{formatted_h}</b></u>", heading_style))
            continue

        # Bulleted lists
        if s.startswith("•") or s.startswith("* "):
            clean_bullet = s.replace("•", "").replace("* ", "").strip()
            formatted_b = format_markdown_for_reportlab(clean_bullet)
            story.append(Paragraph(f"&bull; {formatted_b}", bullet_style))
            continue

        # Regular text
        formatted_text = format_markdown_for_reportlab(s)
        story.append(Paragraph(formatted_text, normal_style))

    doc.build(story)
    output.seek(0)
    return output.getvalue()


# -----------------------------
# -----------------------------
# MAIN WORKFLOW
# -----------------------------
if st.button("🚀 Run", use_container_width=True):
    # Now it checks for the EOB list, MRN, AND the typed CPT code
    if not (eob_files and mrn_file and target_cpt_code_input.strip()):
        st.error("Please upload all required files (at least one EOB and MRN) and enter a Target CPT Code.")
    else:
        with st.spinner("Processing... Please wait..."):
            try:
                # Normalize eob_files to a list if multiple uploads is disabled
                if eob_files and not isinstance(eob_files, list):
                    eob_files = [eob_files]

                mrn_text = extract_text_from_pdf(mrn_file)

                # Initialize variables to hold data
                all_hcpcs_codes = []
                date_of_service, drg_code, billing_provider = "", "", ""

                # Loop through ALL uploaded EOB files
                for i, eob in enumerate(eob_files):
                    eob_text = extract_text_from_pdf(eob)
                    extracted_date, hcpcs_codes, extracted_drg, extracted_billing = extract_fields(eob_text)

                    # 1. Combine all unique codes from EVERY EOB
                    for code in hcpcs_codes:
                        if code not in all_hcpcs_codes:
                            all_hcpcs_codes.append(code)

                    # 2. Grab standard fields from ONLY THE FIRST EOB (Index 0)
                    if i == 0:
                        date_of_service = extracted_date
                        drg_code = extracted_drg
                        billing_provider = extracted_billing

                # --- AUTO-COMMA LOGIC ---
                raw_input = target_cpt_code_input.strip()
                split_codes = re.split(r'[,\s]+', raw_input) 
                target_cpt_code = ", ".join([c for c in split_codes if c])

                # Pass the targeted code to generate the customized summary
                mrn_summary = generate_mrn_summary(mrn_text, target_cpt_code)

                # Generate the final full letter using the COMBINED codes
                full_letter = generate_bcbs_justification_letter(date_of_service, all_hcpcs_codes, drg_code, billing_provider, mrn_summary)

                # Save everything to session state so it doesn't disappear
                st.session_state["original_letter"] = full_letter
                st.session_state["edited_letter"] = full_letter
                st.session_state["doc_generated"] = True

                # --- NEW FILENAME EXTRACTOR ---
                # Grabs the first word from the FIRST uploaded EOB file
                first_word = eob_files[0].name.split()[0].replace(".pdf", "").replace(".PDF", "")
                st.session_state["file_prefix"] = first_word

            except Exception as e:
                st.error(f"❌ Error occurred: {str(e)}")


# -----------------------------
# EDIT & DOWNLOAD SECTION
# -----------------------------

# 1. Define the "revert" logic as a callback function
def revert_to_original():
    st.session_state["edited_letter"] = st.session_state["original_letter"]

# This runs if a document has been successfully generated
if st.session_state.get("doc_generated"):
    llm_info = st.session_state.get("llm_used", "Unknown")
    st.success(f"✅ Automation complete! (Model: {llm_info})")
    st.subheader("📜 Review and Edit BCBS Justification Letter")
    st.caption("You can edit the text directly in the box below. Make sure to click 'Save Edits' before downloading.")

    # 2. Editable Text Box linked to session state
    st.text_area(
        "Make any manual edits below:", 
        key="edited_letter", 
        height=600
    )

    col1, col2 = st.columns(2)
    
    # 3. Save Edits Button
    with col1:
        if st.button("💾 Save Edits", use_container_width=True):
            st.success("Edits saved! The downloaded document will now reflect these changes.")

    # 4. Revert to Original Button (Now using the callback!)
    with col2:
        st.button(
            "🔄 Revert to Original", 
            on_click=revert_to_original, 
            use_container_width=True
        )

    # Generate PDF from the current state of the edited text
    if REPORTLAB_AVAILABLE:
        output_pdf = create_pdf_with_full_letter(st.session_state["edited_letter"])
    else:
        st.error("The 'reportlab' package is required for PDF generation. Please run `pip install reportlab`.")
        output_pdf = b""

    st.markdown("<br>", unsafe_allow_html=True) # Adding a little space

    # Construct the dynamic file name
    prefix = st.session_state.get("file_prefix", "")
    dynamic_file_name = f"{prefix}_BCBS_Justification_for_IDR.pdf" if prefix else "BCBS_Justification_for_IDR.pdf"

    # 5. Download Button
    st.download_button(
        label="📥 Download as PDF (.pdf)",
        data=output_pdf,
        file_name=dynamic_file_name,
        mime="application/pdf",
        use_container_width=True,
        key="download_pdf_button",
        help="Download the generated PDF document (remains available after click)."
    )

